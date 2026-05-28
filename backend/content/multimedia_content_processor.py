"""
KIRO2 Multimedia Content Processing System
Advanced multimedia content processing and optimization system
Türkiye Üniversite Sınavları Hazırlık Platformu - Çoklu Medya İçerik İşleme Sistemi
"""

import asyncio
import mimetypes
import tempfile
import uuid
from dataclasses import dataclass, field
from datetime import UTC, datetime, timedelta
from enum import Enum
from pathlib import Path
from typing import Any

import aiofiles
import cv2
import docx
import librosa
import numpy as np
import soundfile as sf
from moviepy.editor import VideoFileClip
from PIL import Image
from pypdf import PdfReader

from content.unified_content_management import ContentFile
from core.structured_logging import LogCategory, get_logger
from core.unified_config import get_unified_config

logger = get_logger(__name__, LogCategory.CONTENT)
config = get_unified_config()


class ProcessingStatus(Enum):
    """Processing status for multimedia content"""

    PENDING = "pending"
    IN_PROGRESS = "in_progress"
    COMPLETED = "completed"
    FAILED = "failed"
    PARTIAL = "partial"


class MediaFormat(Enum):
    """Supported media formats"""

    # Video formats
    MP4 = "mp4"
    AVI = "avi"
    MOV = "mov"
    WMV = "wmv"
    FLV = "flv"
    WEBM = "webm"

    # Audio formats
    MP3 = "mp3"
    WAV = "wav"
    AAC = "aac"
    OGG = "ogg"
    M4A = "m4a"

    # Image formats
    JPEG = "jpeg"
    PNG = "png"
    GIF = "gif"
    SVG = "svg"
    WEBP = "webp"

    # Document formats
    PDF = "pdf"
    DOCX = "docx"
    XLSX = "xlsx"
    PPTX = "pptx"
    TXT = "txt"


class CompressionLevel(Enum):
    """Compression levels for media optimization"""

    LOW = "low"  # High quality, large file
    MEDIUM = "medium"  # Balanced quality/size
    HIGH = "high"  # Lower quality, small file
    ADAPTIVE = "adaptive"  # AI-based optimization


@dataclass
class ProcessingTask:
    """Task for processing multimedia content"""

    task_id: str
    content_file: ContentFile
    processing_type: str  # compress, convert, extract, analyze
    parameters: dict[str, Any] = field(default_factory=dict)

    # Status and timing
    status: ProcessingStatus = ProcessingStatus.PENDING
    started_at: datetime | None = None
    completed_at: datetime | None = None
    progress_percentage: float = 0.0

    # Results
    output_files: list[ContentFile] = field(default_factory=list)
    extracted_data: dict[str, Any] = field(default_factory=dict)
    processing_errors: list[str] = field(default_factory=list)

    # Resource usage
    cpu_time_seconds: float = 0.0
    memory_peak_mb: float = 0.0

    def __post_init__(self):
        if not self.task_id:
            self.task_id = str(uuid.uuid4())

    def start_processing(self) -> None:
        """Mark task as started"""
        self.status = ProcessingStatus.IN_PROGRESS
        self.started_at = datetime.now(UTC)

    def complete_processing(self, success: bool = True) -> None:
        """Mark task as completed"""
        self.status = ProcessingStatus.COMPLETED if success else ProcessingStatus.FAILED
        self.completed_at = datetime.now(UTC)
        self.progress_percentage = 100.0

    def update_progress(self, percentage: float) -> None:
        """Update processing progress"""
        self.progress_percentage = max(0, min(100, percentage))

    def add_error(self, error: str) -> None:
        """Add processing error"""
        self.processing_errors.append(error)
        logger.error(f"Processing error in task {self.task_id}: {error}")

    def to_dict(self) -> dict[str, Any]:
        """Convert to dictionary"""
        return {
            "task_id": self.task_id,
            "content_file_id": self.content_file.file_id,
            "processing_type": self.processing_type,
            "parameters": self.parameters,
            "status": self.status.value,
            "started_at": self.started_at.isoformat() if self.started_at else None,
            "completed_at": self.completed_at.isoformat()
            if self.completed_at
            else None,
            "progress_percentage": self.progress_percentage,
            "output_files": [f.to_dict() for f in self.output_files],
            "extracted_data": self.extracted_data,
            "processing_errors": self.processing_errors,
            "cpu_time_seconds": self.cpu_time_seconds,
            "memory_peak_mb": self.memory_peak_mb,
        }


@dataclass
class MediaAnalysis:
    """Analysis results for media files"""

    file_id: str
    analysis_type: str

    # Technical properties
    format_info: dict[str, Any] = field(default_factory=dict)
    quality_metrics: dict[str, float] = field(default_factory=dict)

    # Content analysis
    content_features: dict[str, Any] = field(default_factory=dict)
    accessibility_info: dict[str, Any] = field(default_factory=dict)

    # Educational content analysis (for Turkish exam content)
    turkish_content_analysis: dict[str, Any] = field(default_factory=dict)
    detected_text: str | None = None
    detected_language: str = "tr"

    # Performance recommendations
    optimization_suggestions: list[str] = field(default_factory=list)
    compression_recommendations: dict[str, Any] = field(default_factory=dict)

    def to_dict(self) -> dict[str, Any]:
        """Convert to dictionary"""
        return {
            "file_id": self.file_id,
            "analysis_type": self.analysis_type,
            "format_info": self.format_info,
            "quality_metrics": self.quality_metrics,
            "content_features": self.content_features,
            "accessibility_info": self.accessibility_info,
            "turkish_content_analysis": self.turkish_content_analysis,
            "detected_text": self.detected_text,
            "detected_language": self.detected_language,
            "optimization_suggestions": self.optimization_suggestions,
            "compression_recommendations": self.compression_recommendations,
        }


class VideoProcessor:
    """Video processing and optimization"""

    def __init__(self):
        self.supported_formats = [
            MediaFormat.MP4,
            MediaFormat.AVI,
            MediaFormat.MOV,
            MediaFormat.WEBM,
        ]
        self.output_formats = [MediaFormat.MP4, MediaFormat.WEBM]
        self.ffmpeg_path = config.get_setting("media.ffmpeg_path", "ffmpeg")

    async def analyze_video(self, file_path: str) -> MediaAnalysis:
        """Analyze video file properties"""
        analysis = MediaAnalysis(file_id=str(uuid.uuid4()), analysis_type="video")

        try:
            # Use moviepy for basic analysis
            with VideoFileClip(file_path) as video:
                analysis.format_info = {
                    "duration_seconds": video.duration,
                    "fps": video.fps,
                    "width": video.w,
                    "height": video.h,
                    "aspect_ratio": round(video.w / video.h, 2) if video.h > 0 else 0,
                    "has_audio": video.audio is not None,
                }

                # Quality assessment
                analysis.quality_metrics = {
                    "resolution_score": self._calculate_resolution_score(
                        video.w, video.h
                    ),
                    "bitrate_estimation": await self._estimate_video_bitrate(file_path),
                    "frame_rate_score": self._calculate_framerate_score(video.fps),
                }

                # Educational content analysis
                if video.duration > 0:
                    analysis.turkish_content_analysis = {
                        "suitable_for_learning": video.duration
                        <= 1800,  # Max 30 minutes
                        "optimal_duration": 300
                        <= video.duration
                        <= 900,  # 5-15 minutes optimal
                        "attention_span_appropriate": video.duration
                        <= 600,  # 10 minutes max for focus
                        "recommended_segments": max(
                            1, int(video.duration / 300)
                        ),  # 5-minute segments
                    }

                # Generate optimization suggestions
                analysis.optimization_suggestions = (
                    self._generate_video_optimization_suggestions(
                        video.w, video.h, video.duration, video.fps
                    )
                )

        except Exception as e:
            logger.error(f"Error analyzing video {file_path}: {e}")
            analysis.content_features["analysis_error"] = str(e)

        return analysis

    async def compress_video(
        self,
        input_path: str,
        output_path: str,
        compression_level: CompressionLevel = CompressionLevel.MEDIUM,
        target_format: MediaFormat = MediaFormat.MP4,
        progress_callback: callable | None = None,
    ) -> bool:
        """Compress video with specified parameters"""

        try:
            # Define compression parameters based on level
            compression_params = self._get_compression_parameters(
                compression_level, target_format
            )

            # Build ffmpeg command
            cmd = [
                self.ffmpeg_path,
                "-i",
                input_path,
                "-c:v",
                compression_params["video_codec"],
                "-crf",
                str(compression_params["crf"]),
                "-c:a",
                compression_params["audio_codec"],
                "-b:a",
                compression_params["audio_bitrate"],
                "-preset",
                compression_params["preset"],
                "-movflags",
                "+faststart",  # Enable streaming
                "-y",  # Overwrite output
                output_path,
            ]

            # Add resolution scaling if needed
            if compression_params.get("scale"):
                cmd.extend(["-vf", f"scale={compression_params['scale']}"])

            # Execute compression
            process = await asyncio.create_subprocess_exec(
                *cmd, stdout=asyncio.subprocess.PIPE, stderr=asyncio.subprocess.PIPE
            )

            # Monitor progress if callback provided
            if progress_callback:
                await self._monitor_ffmpeg_progress(process, progress_callback)

            stdout, stderr = await process.communicate()

            if process.returncode == 0:
                logger.info(
                    f"Successfully compressed video: {input_path} -> {output_path}"
                )
                return True
            logger.error(f"Video compression failed: {stderr.decode()}")
            return False

        except Exception as e:
            logger.error(f"Error compressing video: {e}")
            return False

    async def extract_thumbnails(
        self,
        video_path: str,
        output_dir: str,
        count: int = 5,
        timestamps: list[float] = None,
    ) -> list[str]:
        """Extract thumbnails from video"""
        thumbnail_paths = []

        try:
            with VideoFileClip(video_path) as video:
                if timestamps is None:
                    # Generate evenly spaced timestamps
                    duration = video.duration
                    timestamps = [
                        duration * i / (count + 1) for i in range(1, count + 1)
                    ]

                for i, timestamp in enumerate(timestamps):
                    if timestamp < video.duration:
                        thumbnail_path = Path(output_dir) / f"thumbnail_{i + 1:03d}.jpg"

                        # Extract frame at timestamp
                        frame = video.get_frame(timestamp)

                        # Convert to PIL Image and save
                        image = Image.fromarray(frame.astype("uint8"), "RGB")
                        image.thumbnail((320, 240), Image.Resampling.LANCZOS)
                        image.save(thumbnail_path, "JPEG", quality=85)

                        thumbnail_paths.append(str(thumbnail_path))

        except Exception as e:
            logger.error(f"Error extracting thumbnails: {e}")

        return thumbnail_paths

    async def extract_audio(self, video_path: str, audio_output_path: str) -> bool:
        """Extract audio from video"""
        try:
            cmd = [
                self.ffmpeg_path,
                "-i",
                video_path,
                "-vn",  # No video
                "-acodec",
                "mp3",
                "-ab",
                "128k",
                "-ar",
                "44100",
                "-y",
                audio_output_path,
            ]

            process = await asyncio.create_subprocess_exec(
                *cmd, stdout=asyncio.subprocess.PIPE, stderr=asyncio.subprocess.PIPE
            )

            stdout, stderr = await process.communicate()

            if process.returncode == 0:
                logger.info(
                    f"Successfully extracted audio: {video_path} -> {audio_output_path}"
                )
                return True
            logger.error(f"Audio extraction failed: {stderr.decode()}")
            return False

        except Exception as e:
            logger.error(f"Error extracting audio: {e}")
            return False

    def _calculate_resolution_score(self, width: int, height: int) -> float:
        """Calculate resolution quality score (0-100)"""
        pixel_count = width * height

        # Score based on common resolutions
        if pixel_count >= 3840 * 2160:  # 4K
            return 100
        if pixel_count >= 1920 * 1080:  # Full HD
            return 90
        if pixel_count >= 1280 * 720:  # HD
            return 75
        if pixel_count >= 854 * 480:  # 480p
            return 60
        if pixel_count >= 640 * 360:  # 360p
            return 40
        return 20

    def _calculate_framerate_score(self, fps: float) -> float:
        """Calculate frame rate quality score (0-100)"""
        if fps >= 60:
            return 100
        if fps >= 30:
            return 90
        if fps >= 24:
            return 75
        if fps >= 15:
            return 50
        return 25

    async def _estimate_video_bitrate(self, file_path: str) -> float:
        """Estimate video bitrate"""
        try:
            file_size = Path(file_path).stat().st_size
            with VideoFileClip(file_path) as video:
                duration = video.duration
                if duration > 0:
                    return (file_size * 8) / (duration * 1024 * 1024)  # Mbps
        except Exception as e:
            logger.error(f"Error estimating bitrate: {e}")
        return 0.0

    def _get_compression_parameters(
        self, compression_level: CompressionLevel, target_format: MediaFormat
    ) -> dict[str, Any]:
        """Get compression parameters for given level and format"""
        base_params = {
            MediaFormat.MP4: {
                "video_codec": "libx264",
                "audio_codec": "aac",
                "preset": "medium",
            },
            MediaFormat.WEBM: {
                "video_codec": "libvpx-vp9",
                "audio_codec": "libopus",
                "preset": "medium",
            },
        }

        params = base_params.get(target_format, base_params[MediaFormat.MP4])

        # Adjust parameters based on compression level
        if compression_level == CompressionLevel.LOW:
            params.update({"crf": 18, "audio_bitrate": "192k", "preset": "slower"})
        elif compression_level == CompressionLevel.MEDIUM:
            params.update({"crf": 23, "audio_bitrate": "128k", "preset": "medium"})
        elif compression_level == CompressionLevel.HIGH:
            params.update(
                {
                    "crf": 28,
                    "audio_bitrate": "96k",
                    "preset": "fast",
                    "scale": "1280:720",  # Downscale to 720p
                }
            )
        else:  # ADAPTIVE
            # AI-based optimization would go here
            params.update({"crf": 25, "audio_bitrate": "128k", "preset": "medium"})

        return params

    def _generate_video_optimization_suggestions(
        self, width: int, height: int, duration: float, fps: float
    ) -> list[str]:
        """Generate optimization suggestions for video"""
        suggestions = []

        # Resolution suggestions
        if width > 1920 or height > 1080:
            suggestions.append("Aşırı yüksek çözünürlük - 1080p'ye düşürmeyi düşünün")
        elif width > 1280 or height > 720:
            suggestions.append("Mobil cihazlar için 720p versiyonu oluşturun")

        # Frame rate suggestions
        if fps > 30:
            suggestions.append("30 FPS'ye düşürerek dosya boyutunu küçültebilirsiniz")
        elif fps < 24:
            suggestions.append("En az 24 FPS kullanarak görüntü akıcılığını artırın")

        # Duration suggestions
        if duration > 1800:  # 30 minutes
            suggestions.append("Uzun videolar - bölümlere ayırmayı düşünün")
        elif duration > 900:  # 15 minutes
            suggestions.append(
                "Öğrenci dikkatini korumak için 10-15 dakikalık segmentler oluşturun"
            )

        # Educational content suggestions
        suggestions.extend(
            [
                "Video başına 1-2 önemli konsepte odaklanın",
                "Her 5 dakikada bir özet veya checkpoint ekleyin",
                "Türkçe altyazı eklemeyi düşünün",
            ]
        )

        return suggestions

    async def _monitor_ffmpeg_progress(self, process, progress_callback):
        """Monitor FFmpeg progress and call callback"""
        # This is a simplified progress monitor
        # In reality, you'd parse FFmpeg's progress output
        for progress in range(0, 101, 10):
            if process.returncode is not None:
                break
            await progress_callback(progress)
            await asyncio.sleep(1)


class AudioProcessor:
    """Audio processing and optimization"""

    def __init__(self):
        self.supported_formats = [
            MediaFormat.MP3,
            MediaFormat.WAV,
            MediaFormat.AAC,
            MediaFormat.OGG,
        ]
        self.output_formats = [MediaFormat.MP3, MediaFormat.AAC, MediaFormat.OGG]

    async def analyze_audio(self, file_path: str) -> MediaAnalysis:
        """Analyze audio file properties"""
        analysis = MediaAnalysis(file_id=str(uuid.uuid4()), analysis_type="audio")

        try:
            # Load audio file
            y, sr = librosa.load(file_path)
            duration = librosa.get_duration(y=y, sr=sr)

            analysis.format_info = {
                "duration_seconds": duration,
                "sample_rate": sr,
                "channels": 1 if y.ndim == 1 else y.shape[0],
                "samples": len(y),
            }

            # Audio quality analysis
            analysis.quality_metrics = {
                "rms_energy": float(np.sqrt(np.mean(y**2))),
                "zero_crossing_rate": float(
                    np.mean(librosa.feature.zero_crossing_rate(y))
                ),
                "spectral_centroid": float(
                    np.mean(librosa.feature.spectral_centroid(y=y, sr=sr))
                ),
                "dynamic_range": float(np.max(y) - np.min(y)),
            }

            # Detect silence and speech
            intervals = librosa.effects.split(y, top_db=20)
            speech_duration = sum((end - start) / sr for start, end in intervals)

            analysis.content_features = {
                "speech_ratio": speech_duration / duration if duration > 0 else 0,
                "silence_duration": duration - speech_duration,
                "speech_segments": len(intervals),
                "average_segment_length": speech_duration / len(intervals)
                if intervals
                else 0,
            }

            # Educational content analysis
            analysis.turkish_content_analysis = {
                "suitable_for_listening": duration <= 3600,  # Max 1 hour
                "optimal_duration": 180 <= duration <= 1200,  # 3-20 minutes optimal
                "clear_speech_detected": analysis.content_features["speech_ratio"]
                > 0.7,
                "background_noise_level": "low"
                if analysis.quality_metrics["rms_energy"] < 0.1
                else "medium",
            }

            # Generate suggestions
            analysis.optimization_suggestions = (
                self._generate_audio_optimization_suggestions(
                    duration, analysis.quality_metrics, analysis.content_features
                )
            )

        except Exception as e:
            logger.error(f"Error analyzing audio {file_path}: {e}")
            analysis.content_features["analysis_error"] = str(e)

        return analysis

    async def compress_audio(
        self,
        input_path: str,
        output_path: str,
        compression_level: CompressionLevel = CompressionLevel.MEDIUM,
        target_format: MediaFormat = MediaFormat.MP3,
        progress_callback: callable | None = None,
    ) -> bool:
        """Compress audio file"""
        try:
            # Load audio
            y, sr = librosa.load(input_path)

            # Apply compression based on level
            if compression_level == CompressionLevel.HIGH:
                # Reduce sample rate for high compression
                sr = min(sr, 22050)
                y = librosa.resample(y, orig_sr=sr, target_sr=22050)
                sr = 22050

            # Normalize audio
            y = librosa.util.normalize(y)

            # Save compressed audio
            if target_format == MediaFormat.MP3:
                # Use pydub for MP3 export (requires ffmpeg)
                import tempfile

                temp_wav = tempfile.mktemp(suffix=".wav")
                sf.write(temp_wav, y, sr)

                # Convert to MP3 using ffmpeg
                bitrate = self._get_audio_bitrate(compression_level)
                cmd = [
                    "ffmpeg",
                    "-i",
                    temp_wav,
                    "-acodec",
                    "mp3",
                    "-ab",
                    bitrate,
                    "-y",
                    output_path,
                ]

                process = await asyncio.create_subprocess_exec(
                    *cmd, stdout=asyncio.subprocess.PIPE, stderr=asyncio.subprocess.PIPE
                )

                stdout, stderr = await process.communicate()
                Path(temp_wav).unlink()  # Clean up temp file

                success = process.returncode == 0
            else:
                # Use soundfile for other formats
                sf.write(output_path, y, sr)
                success = True

            if success:
                logger.info(
                    f"Successfully compressed audio: {input_path} -> {output_path}"
                )
            else:
                logger.error("Audio compression failed")

            return success

        except Exception as e:
            logger.error(f"Error compressing audio: {e}")
            return False

    def _get_audio_bitrate(self, compression_level: CompressionLevel) -> str:
        """Get audio bitrate for compression level"""
        bitrates = {
            CompressionLevel.LOW: "320k",
            CompressionLevel.MEDIUM: "192k",
            CompressionLevel.HIGH: "128k",
            CompressionLevel.ADAPTIVE: "192k",
        }
        return bitrates.get(compression_level, "192k")

    def _generate_audio_optimization_suggestions(
        self,
        duration: float,
        quality_metrics: dict[str, float],
        content_features: dict[str, Any],
    ) -> list[str]:
        """Generate optimization suggestions for audio"""
        suggestions = []

        # Duration suggestions
        if duration > 3600:  # 1 hour
            suggestions.append("Uzun ses kayıtları - bölümlere ayırın")
        elif duration > 1800:  # 30 minutes
            suggestions.append(
                "30 dakikadan uzun - öğrenci dikkatini korumak için segmentlere ayırın"
            )

        # Quality suggestions
        if quality_metrics.get("rms_energy", 0) < 0.01:
            suggestions.append("Ses seviyesi çok düşük - normalize edin")
        elif quality_metrics.get("rms_energy", 0) > 0.5:
            suggestions.append("Ses seviyesi çok yüksek - ses limitleyici kullanın")

        # Content suggestions
        speech_ratio = content_features.get("speech_ratio", 0)
        if speech_ratio < 0.5:
            suggestions.append("Konuşma oranı düşük - sessizlikleri kaldırın")

        # Educational suggestions
        suggestions.extend(
            [
                "Önemli noktaları vurgulamak için ses tonunu değiştirin",
                "Anlayışı artırmak için yavaş ve net konuşun",
                "Her bölüm için özet ekleyin",
            ]
        )

        return suggestions


class ImageProcessor:
    """Image processing and optimization"""

    def __init__(self):
        self.supported_formats = [
            MediaFormat.JPEG,
            MediaFormat.PNG,
            MediaFormat.GIF,
            MediaFormat.WEBP,
        ]
        self.output_formats = [MediaFormat.JPEG, MediaFormat.PNG, MediaFormat.WEBP]

    async def analyze_image(self, file_path: str) -> MediaAnalysis:
        """Analyze image properties"""
        analysis = MediaAnalysis(file_id=str(uuid.uuid4()), analysis_type="image")

        try:
            with Image.open(file_path) as img:
                analysis.format_info = {
                    "width": img.width,
                    "height": img.height,
                    "mode": img.mode,
                    "format": img.format,
                    "aspect_ratio": round(img.width / img.height, 2)
                    if img.height > 0
                    else 0,
                    "has_transparency": img.mode in ["RGBA", "LA", "P"],
                }

                # Quality assessment
                file_size = Path(file_path).stat().st_size
                pixel_count = img.width * img.height

                analysis.quality_metrics = {
                    "resolution_score": self._calculate_image_resolution_score(
                        img.width, img.height
                    ),
                    "file_efficiency": pixel_count / file_size if file_size > 0 else 0,
                    "compression_ratio": file_size / pixel_count
                    if pixel_count > 0
                    else 0,
                }

                # Detect if image contains text (OCR would go here)
                analysis.content_features = {
                    "likely_contains_text": self._detect_text_regions(img),
                    "dominant_colors": self._extract_dominant_colors(img),
                    "brightness": self._calculate_brightness(img),
                    "contrast": self._calculate_contrast(img),
                }

                # Educational content analysis
                analysis.turkish_content_analysis = {
                    "suitable_for_web": img.width <= 1920 and img.height <= 1080,
                    "mobile_friendly": img.width <= 800 and img.height <= 600,
                    "diagram_suitable": analysis.content_features[
                        "likely_contains_text"
                    ],
                    "accessibility_friendly": analysis.content_features["contrast"]
                    > 0.3,
                }

                # Generate suggestions
                analysis.optimization_suggestions = (
                    self._generate_image_optimization_suggestions(
                        img.width, img.height, file_size, img.format
                    )
                )

        except Exception as e:
            logger.error(f"Error analyzing image {file_path}: {e}")
            analysis.content_features["analysis_error"] = str(e)

        return analysis

    async def compress_image(
        self,
        input_path: str,
        output_path: str,
        compression_level: CompressionLevel = CompressionLevel.MEDIUM,
        target_format: MediaFormat = MediaFormat.JPEG,
        max_width: int = 1920,
        max_height: int = 1080,
    ) -> bool:
        """Compress and optimize image"""
        try:
            with Image.open(input_path) as img:
                # Convert RGBA to RGB if saving as JPEG
                if target_format == MediaFormat.JPEG and img.mode in [
                    "RGBA",
                    "LA",
                    "P",
                ]:
                    if img.mode == "P":
                        img = img.convert("RGBA")

                    # Create white background
                    background = Image.new("RGB", img.size, (255, 255, 255))
                    background.paste(
                        img, mask=img.split()[-1] if "A" in img.mode else None
                    )
                    img = background

                # Resize if necessary
                if img.width > max_width or img.height > max_height:
                    img.thumbnail((max_width, max_height), Image.Resampling.LANCZOS)

                # Set quality based on compression level
                quality_map = {
                    CompressionLevel.LOW: 95,
                    CompressionLevel.MEDIUM: 85,
                    CompressionLevel.HIGH: 75,
                    CompressionLevel.ADAPTIVE: 80,
                }
                quality = quality_map.get(compression_level, 85)

                # Save optimized image
                save_kwargs = {"optimize": True}
                if target_format == MediaFormat.JPEG:
                    save_kwargs["quality"] = quality
                    save_kwargs["progressive"] = True
                elif target_format == MediaFormat.PNG:
                    save_kwargs["compress_level"] = 6
                elif target_format == MediaFormat.WEBP:
                    save_kwargs["quality"] = quality
                    save_kwargs["method"] = 6

                img.save(output_path, format=target_format.value.upper(), **save_kwargs)

            logger.info(f"Successfully compressed image: {input_path} -> {output_path}")
            return True

        except Exception as e:
            logger.error(f"Error compressing image: {e}")
            return False

    async def generate_thumbnails(
        self, input_path: str, output_dir: str, sizes: list[tuple[int, int]] = None
    ) -> list[str]:
        """Generate multiple thumbnail sizes"""
        if sizes is None:
            sizes = [(150, 150), (300, 300), (600, 400)]

        thumbnail_paths = []

        try:
            with Image.open(input_path) as img:
                for i, (width, height) in enumerate(sizes):
                    thumbnail = img.copy()
                    thumbnail.thumbnail((width, height), Image.Resampling.LANCZOS)

                    # Maintain aspect ratio by adding padding if needed
                    if thumbnail.size != (width, height):
                        padded = Image.new("RGB", (width, height), (255, 255, 255))
                        paste_x = (width - thumbnail.width) // 2
                        paste_y = (height - thumbnail.height) // 2
                        padded.paste(thumbnail, (paste_x, paste_y))
                        thumbnail = padded

                    thumbnail_path = Path(output_dir) / f"thumb_{width}x{height}.jpg"
                    thumbnail.save(thumbnail_path, "JPEG", quality=85, optimize=True)
                    thumbnail_paths.append(str(thumbnail_path))

        except Exception as e:
            logger.error(f"Error generating thumbnails: {e}")

        return thumbnail_paths

    def _calculate_image_resolution_score(self, width: int, height: int) -> float:
        """Calculate image resolution quality score"""
        pixel_count = width * height

        if pixel_count >= 1920 * 1080:  # Full HD
            return 100
        if pixel_count >= 1280 * 720:  # HD
            return 85
        if pixel_count >= 800 * 600:  # SVGA
            return 70
        if pixel_count >= 640 * 480:  # VGA
            return 55
        return 30

    def _detect_text_regions(self, img: Image.Image) -> bool:
        """Detect if image likely contains text (simplified)"""
        # Convert to grayscale
        gray = img.convert("L")
        np_img = np.array(gray)

        # Simple edge detection to find text-like regions
        # In production, would use proper OCR
        edges = cv2.Canny(np_img, 50, 150)
        edge_ratio = np.sum(edges > 0) / edges.size

        return edge_ratio > 0.05  # Threshold for text detection

    def _extract_dominant_colors(self, img: Image.Image, k: int = 5) -> list[list[int]]:
        """Extract dominant colors from image"""
        # Resize for faster processing
        img_small = img.resize((150, 150))
        np_img = np.array(img_small)

        if len(np_img.shape) == 3:
            # Reshape for K-means
            data = np_img.reshape((-1, 3))
            data = np.float32(data)

            # Apply K-means clustering
            criteria = (cv2.TERM_CRITERIA_EPS + cv2.TERM_CRITERIA_MAX_ITER, 20, 1.0)
            _, labels, centers = cv2.kmeans(
                data, k, None, criteria, 10, cv2.KMEANS_RANDOM_CENTERS
            )

            # Convert centers back to int and return
            centers = np.uint8(centers)
            return [center.tolist() for center in centers]

        return []

    def _calculate_brightness(self, img: Image.Image) -> float:
        """Calculate average brightness of image"""
        gray = img.convert("L")
        np_img = np.array(gray)
        return float(np.mean(np_img) / 255.0)

    def _calculate_contrast(self, img: Image.Image) -> float:
        """Calculate contrast of image"""
        gray = img.convert("L")
        np_img = np.array(gray)
        return float(np.std(np_img) / 255.0)

    def _generate_image_optimization_suggestions(
        self, width: int, height: int, file_size: int, format: str
    ) -> list[str]:
        """Generate optimization suggestions for image"""
        suggestions = []

        # Resolution suggestions
        if width > 1920 or height > 1080:
            suggestions.append("Yüksek çözünürlük - web kullanımı için küçültün")

        # Format suggestions
        if format == "PNG" and file_size > 500000:  # 500KB
            suggestions.append("Büyük PNG dosyası - JPEG formatını düşünün")
        elif format == "BMP":
            suggestions.append("BMP formatı verimsiz - JPEG veya PNG kullanın")

        # Size suggestions
        if file_size > 2000000:  # 2MB
            suggestions.append("Büyük dosya boyutu - sıkıştırma uygulayın")

        # Educational content suggestions
        suggestions.extend(
            [
                "Diyagramlar için yüksek kontrast kullanın",
                "Mobil uyumlu boyutlarda thumbnail oluşturun",
                "Önemli metinlerin okunabilir olduğundan emin olun",
            ]
        )

        return suggestions


class DocumentProcessor:
    """Document processing and text extraction"""

    def __init__(self):
        self.supported_formats = [MediaFormat.PDF, MediaFormat.DOCX, MediaFormat.TXT]

    async def analyze_document(self, file_path: str) -> MediaAnalysis:
        """Analyze document properties and extract text"""
        analysis = MediaAnalysis(file_id=str(uuid.uuid4()), analysis_type="document")

        file_path_obj = Path(file_path)
        file_ext = file_path_obj.suffix.lower()

        try:
            if file_ext == ".pdf":
                analysis = await self._analyze_pdf(file_path, analysis)
            elif file_ext == ".docx":
                analysis = await self._analyze_docx(file_path, analysis)
            elif file_ext == ".txt":
                analysis = await self._analyze_txt(file_path, analysis)
            else:
                analysis.content_features["analysis_error"] = (
                    f"Unsupported format: {file_ext}"
                )

        except Exception as e:
            logger.error(f"Error analyzing document {file_path}: {e}")
            analysis.content_features["analysis_error"] = str(e)

        return analysis

    async def _analyze_pdf(
        self, file_path: str, analysis: MediaAnalysis
    ) -> MediaAnalysis:
        """Analyze PDF document"""
        with open(file_path, "rb") as file:
            reader = PdfReader(file)

            analysis.format_info = {
                "page_count": len(reader.pages),
                "has_text": True,
                "has_images": False,  # Would need deeper inspection
                "encrypted": reader.is_encrypted,
            }

            # Extract text from all pages
            full_text = ""
            for page in reader.pages:
                full_text += page.extract_text() + "\n"

            analysis.detected_text = full_text.strip()

            # Analyze text content
            word_count = len(full_text.split())
            char_count = len(full_text)

            analysis.content_features = {
                "word_count": word_count,
                "character_count": char_count,
                "average_words_per_page": word_count / len(reader.pages)
                if reader.pages
                else 0,
                "reading_time_minutes": word_count / 200
                if word_count > 0
                else 0,  # Assume 200 WPM
            }

            # Turkish content analysis
            analysis.turkish_content_analysis = {
                "suitable_length": 500
                <= word_count
                <= 5000,  # Reasonable length for education
                "contains_turkish_text": self._detect_turkish_content(full_text),
                "readability_appropriate": True,  # Would use Turkish readability metrics
                "academic_content": self._detect_academic_keywords(full_text),
            }

            # Generate suggestions
            analysis.optimization_suggestions = self._generate_document_suggestions(
                len(reader.pages), word_count, analysis.turkish_content_analysis
            )

        return analysis

    async def _analyze_docx(
        self, file_path: str, analysis: MediaAnalysis
    ) -> MediaAnalysis:
        """Analyze DOCX document"""
        doc = docx.Document(file_path)

        # Extract text
        full_text = ""
        paragraph_count = 0

        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                full_text += paragraph.text + "\n"
                paragraph_count += 1

        analysis.detected_text = full_text.strip()

        # Count tables and images
        table_count = len(doc.tables)
        image_count = 0

        for rel in doc.part.rels.values():
            if "image" in rel.target_ref:
                image_count += 1

        analysis.format_info = {
            "paragraph_count": paragraph_count,
            "table_count": table_count,
            "image_count": image_count,
            "has_text": bool(full_text.strip()),
            "has_formatting": True,
        }

        # Analyze content
        word_count = len(full_text.split())

        analysis.content_features = {
            "word_count": word_count,
            "character_count": len(full_text),
            "reading_time_minutes": word_count / 200 if word_count > 0 else 0,
            "structural_elements": table_count + image_count,
        }

        # Turkish content analysis
        analysis.turkish_content_analysis = {
            "suitable_length": 300 <= word_count <= 3000,
            "contains_turkish_text": self._detect_turkish_content(full_text),
            "well_structured": paragraph_count > 3,
            "has_multimedia": image_count > 0 or table_count > 0,
        }

        return analysis

    async def _analyze_txt(
        self, file_path: str, analysis: MediaAnalysis
    ) -> MediaAnalysis:
        """Analyze plain text document"""
        async with aiofiles.open(file_path, encoding="utf-8") as file:
            full_text = await file.read()

        analysis.detected_text = full_text

        lines = full_text.split("\n")
        words = full_text.split()

        analysis.format_info = {
            "line_count": len(lines),
            "encoding": "utf-8",
            "has_text": bool(full_text.strip()),
        }

        analysis.content_features = {
            "word_count": len(words),
            "character_count": len(full_text),
            "line_count": len([line for line in lines if line.strip()]),
            "reading_time_minutes": len(words) / 200 if words else 0,
        }

        # Turkish content analysis
        analysis.turkish_content_analysis = {
            "suitable_length": 200 <= len(words) <= 2000,
            "contains_turkish_text": self._detect_turkish_content(full_text),
            "plain_text_appropriate": True,
        }

        return analysis

    def _detect_turkish_content(self, text: str) -> bool:
        """Detect if text contains Turkish content"""
        # Simple Turkish character detection
        turkish_chars = set("çğıöşüÇĞIİÖŞÜ")
        text_chars = set(text)
        return len(turkish_chars.intersection(text_chars)) > 0

    def _detect_academic_keywords(self, text: str) -> bool:
        """Detect if document contains academic/educational keywords"""
        academic_keywords = [
            "soru",
            "cevap",
            "çözüm",
            "örnek",
            "konu",
            "ders",
            "bölüm",
            "matematik",
            "fizik",
            "kimya",
            "biyoloji",
            "türkçe",
            "tarih",
            "tyt",
            "ayt",
            "yks",
            "sınav",
            "test",
            "deneme",
        ]

        text_lower = text.lower()
        return any(keyword in text_lower for keyword in academic_keywords)

    def _generate_document_suggestions(
        self, page_count: int, word_count: int, turkish_analysis: dict[str, Any]
    ) -> list[str]:
        """Generate optimization suggestions for document"""
        suggestions = []

        # Length suggestions
        if word_count > 5000:
            suggestions.append("Uzun belge - bölümlere ayırın")
        elif word_count < 200:
            suggestions.append("Çok kısa belge - daha detaylı açıklama ekleyin")

        # Structure suggestions
        if page_count > 20:
            suggestions.append("Çok sayfa - özet ve içindekiler ekleyin")

        # Content suggestions
        if not turkish_analysis.get("contains_turkish_text", False):
            suggestions.append("Türkçe içerik tespit edilmedi")

        # Educational suggestions
        suggestions.extend(
            [
                "Önemli kavramları vurgulayın",
                "Örnekler ve alıştırmalar ekleyin",
                "Her bölümün sonuna özet ekleyin",
                "Görsel destekli açıklamalar kullanın",
            ]
        )

        return suggestions


class MultimediaContentProcessor:
    """Main multimedia content processing orchestrator"""

    def __init__(self):
        self.video_processor = VideoProcessor()
        self.audio_processor = AudioProcessor()
        self.image_processor = ImageProcessor()
        self.document_processor = DocumentProcessor()

        # Processing queues
        self.processing_queue: asyncio.Queue = asyncio.Queue()
        self.active_tasks: dict[str, ProcessingTask] = {}
        self.completed_tasks: dict[str, ProcessingTask] = {}

        # Configuration
        self.max_concurrent_tasks = config.get_setting("media.max_concurrent_tasks", 4)
        self.temp_dir = Path(
            config.get_setting("media.temp_dir", tempfile.gettempdir())
        )
        self.output_dir = Path(config.get_setting("media.output_dir", "./media_output"))

        # Ensure directories exist
        self.temp_dir.mkdir(exist_ok=True)
        self.output_dir.mkdir(exist_ok=True)

    async def start_processing_service(self) -> None:
        """Start the multimedia processing service"""
        logger.info("Starting multimedia content processing service")

        # Start worker tasks
        workers = []
        for i in range(self.max_concurrent_tasks):
            worker = asyncio.create_task(self._processing_worker(f"worker-{i}"))
            workers.append(worker)

        # Wait for all workers
        await asyncio.gather(*workers)

    async def _processing_worker(self, worker_name: str) -> None:
        """Worker task for processing multimedia content"""
        logger.info(f"Started processing worker: {worker_name}")

        while True:
            try:
                # Get next task from queue
                task = await self.processing_queue.get()

                logger.info(f"{worker_name} processing task {task.task_id}")

                # Move task to active tasks
                self.active_tasks[task.task_id] = task
                task.start_processing()

                # Process based on type
                success = await self._process_task(task)

                # Complete task
                task.complete_processing(success)

                # Move to completed tasks
                del self.active_tasks[task.task_id]
                self.completed_tasks[task.task_id] = task

                # Mark queue task as done
                self.processing_queue.task_done()

            except Exception as e:
                logger.error(f"Error in processing worker {worker_name}: {e}")
                if "task" in locals():
                    task.add_error(str(e))
                    task.complete_processing(False)

    async def submit_processing_task(self, task: ProcessingTask) -> str:
        """Submit a processing task to the queue"""
        await self.processing_queue.put(task)
        logger.info(
            f"Submitted processing task {task.task_id} of type {task.processing_type}"
        )
        return task.task_id

    async def _process_task(self, task: ProcessingTask) -> bool:
        """Process individual task based on its type"""
        try:
            file_path = task.content_file.file_path

            if task.processing_type == "analyze":
                return await self._analyze_content(task)
            if task.processing_type == "compress":
                return await self._compress_content(task)
            if task.processing_type == "convert":
                return await self._convert_content(task)
            if task.processing_type == "extract":
                return await self._extract_content_features(task)
            task.add_error(f"Unknown processing type: {task.processing_type}")
            return False

        except Exception as e:
            task.add_error(f"Processing failed: {e!s}")
            return False

    async def _analyze_content(self, task: ProcessingTask) -> bool:
        """Analyze content and extract metadata"""
        file_path = task.content_file.file_path
        mime_type = task.content_file.mime_type

        try:
            analysis = None

            if mime_type.startswith("video/"):
                analysis = await self.video_processor.analyze_video(file_path)
            elif mime_type.startswith("audio/"):
                analysis = await self.audio_processor.analyze_audio(file_path)
            elif mime_type.startswith("image/"):
                analysis = await self.image_processor.analyze_image(file_path)
            elif mime_type in ["application/pdf", "application/msword", "text/plain"]:
                analysis = await self.document_processor.analyze_document(file_path)
            else:
                task.add_error(f"Unsupported MIME type for analysis: {mime_type}")
                return False

            if analysis:
                task.extracted_data["analysis"] = analysis.to_dict()
                return True

        except Exception as e:
            task.add_error(f"Analysis failed: {e!s}")

        return False

    async def _compress_content(self, task: ProcessingTask) -> bool:
        """Compress content based on type"""
        file_path = task.content_file.file_path
        mime_type = task.content_file.mime_type

        # Get compression parameters
        compression_level = CompressionLevel(
            task.parameters.get("compression_level", "medium")
        )

        try:
            # Generate output path
            input_path = Path(file_path)
            output_path = (
                self.output_dir / f"compressed_{uuid.uuid4()}{input_path.suffix}"
            )

            success = False

            if mime_type.startswith("video/"):
                target_format = MediaFormat(task.parameters.get("target_format", "mp4"))
                success = await self.video_processor.compress_video(
                    str(input_path), str(output_path), compression_level, target_format
                )
            elif mime_type.startswith("audio/"):
                target_format = MediaFormat(task.parameters.get("target_format", "mp3"))
                success = await self.audio_processor.compress_audio(
                    str(input_path), str(output_path), compression_level, target_format
                )
            elif mime_type.startswith("image/"):
                target_format = MediaFormat(
                    task.parameters.get("target_format", "jpeg")
                )
                success = await self.image_processor.compress_image(
                    str(input_path), str(output_path), compression_level, target_format
                )

            if success and output_path.exists():
                # Create ContentFile for output
                output_file = ContentFile(
                    file_id=str(uuid.uuid4()),
                    filename=output_path.name,
                    file_path=str(output_path),
                    file_size=output_path.stat().st_size,
                    mime_type=mimetypes.guess_type(str(output_path))[0] or mime_type,
                    upload_date=datetime.now(UTC),
                    uploaded_by=task.content_file.uploaded_by,
                    processed=True,
                    processing_status="compressed",
                )

                task.output_files.append(output_file)
                return True

        except Exception as e:
            task.add_error(f"Compression failed: {e!s}")

        return False

    async def _convert_content(self, task: ProcessingTask) -> bool:
        """Convert content to different format"""
        # Implementation would be similar to compression but with format conversion
        task.add_error("Format conversion not yet implemented")
        return False

    async def _extract_content_features(self, task: ProcessingTask) -> bool:
        """Extract specific features from content"""
        file_path = task.content_file.file_path
        mime_type = task.content_file.mime_type
        feature_type = task.parameters.get("feature_type", "thumbnails")

        try:
            if feature_type == "thumbnails" and mime_type.startswith("video/"):
                # Extract video thumbnails
                thumb_dir = self.temp_dir / f"thumbnails_{uuid.uuid4()}"
                thumb_dir.mkdir(exist_ok=True)

                thumbnail_paths = await self.video_processor.extract_thumbnails(
                    file_path, str(thumb_dir), count=5
                )

                # Create ContentFiles for thumbnails
                for thumb_path in thumbnail_paths:
                    thumb_file = ContentFile(
                        file_id=str(uuid.uuid4()),
                        filename=Path(thumb_path).name,
                        file_path=thumb_path,
                        file_size=Path(thumb_path).stat().st_size,
                        mime_type="image/jpeg",
                        upload_date=datetime.now(UTC),
                        uploaded_by=task.content_file.uploaded_by,
                        processed=True,
                        processing_status="thumbnail",
                    )
                    task.output_files.append(thumb_file)

                return len(thumbnail_paths) > 0

            if feature_type == "audio" and mime_type.startswith("video/"):
                # Extract audio from video
                audio_path = self.temp_dir / f"audio_{uuid.uuid4()}.mp3"

                success = await self.video_processor.extract_audio(
                    file_path, str(audio_path)
                )

                if success and audio_path.exists():
                    audio_file = ContentFile(
                        file_id=str(uuid.uuid4()),
                        filename=audio_path.name,
                        file_path=str(audio_path),
                        file_size=audio_path.stat().st_size,
                        mime_type="audio/mpeg",
                        upload_date=datetime.now(UTC),
                        uploaded_by=task.content_file.uploaded_by,
                        processed=True,
                        processing_status="extracted_audio",
                    )
                    task.output_files.append(audio_file)
                    return True

            elif feature_type == "thumbnails" and mime_type.startswith("image/"):
                # Generate image thumbnails
                thumb_dir = self.temp_dir / f"thumbnails_{uuid.uuid4()}"
                thumb_dir.mkdir(exist_ok=True)

                thumbnail_paths = await self.image_processor.generate_thumbnails(
                    file_path, str(thumb_dir)
                )

                for thumb_path in thumbnail_paths:
                    thumb_file = ContentFile(
                        file_id=str(uuid.uuid4()),
                        filename=Path(thumb_path).name,
                        file_path=thumb_path,
                        file_size=Path(thumb_path).stat().st_size,
                        mime_type="image/jpeg",
                        upload_date=datetime.now(UTC),
                        uploaded_by=task.content_file.uploaded_by,
                        processed=True,
                        processing_status="thumbnail",
                    )
                    task.output_files.append(thumb_file)

                return len(thumbnail_paths) > 0

        except Exception as e:
            task.add_error(f"Feature extraction failed: {e!s}")

        return False

    async def get_task_status(self, task_id: str) -> dict[str, Any] | None:
        """Get status of processing task"""
        if task_id in self.active_tasks:
            return self.active_tasks[task_id].to_dict()
        if task_id in self.completed_tasks:
            return self.completed_tasks[task_id].to_dict()
        return None

    async def get_processing_statistics(self) -> dict[str, Any]:
        """Get processing service statistics"""
        return {
            "queue_size": self.processing_queue.qsize(),
            "active_tasks": len(self.active_tasks),
            "completed_tasks": len(self.completed_tasks),
            "max_concurrent_tasks": self.max_concurrent_tasks,
            "temp_dir_size_mb": sum(
                f.stat().st_size for f in self.temp_dir.rglob("*") if f.is_file()
            )
            / (1024 * 1024),
            "active_task_types": [
                task.processing_type for task in self.active_tasks.values()
            ],
        }

    async def cleanup_temp_files(self, older_than_hours: int = 24) -> int:
        """Clean up temporary files older than specified hours"""
        cutoff_time = datetime.now() - timedelta(hours=older_than_hours)
        cleaned_count = 0

        try:
            for file_path in self.temp_dir.rglob("*"):
                if file_path.is_file():
                    file_time = datetime.fromtimestamp(file_path.stat().st_mtime)
                    if file_time < cutoff_time:
                        file_path.unlink()
                        cleaned_count += 1

            logger.info(f"Cleaned up {cleaned_count} temporary files")
        except Exception as e:
            logger.error(f"Error cleaning temporary files: {e}")

        return cleaned_count


if __name__ == "__main__":
    # Example usage and testing
    print("KIRO2 Multimedia Content Processing System")
    print("=" * 45)

    async def test_processing_system():
        """Test multimedia processing system"""
        processor = MultimediaContentProcessor()

        # Create sample content file (would normally be uploaded)
        sample_file = ContentFile(
            file_id=str(uuid.uuid4()),
            filename="sample_video.mp4",
            file_path="/path/to/sample/video.mp4",  # Would be real path
            file_size=10000000,  # 10MB
            mime_type="video/mp4",
            upload_date=datetime.now(UTC),
            uploaded_by=1001,
        )

        # Create analysis task
        analysis_task = ProcessingTask(
            task_id=str(uuid.uuid4()),
            content_file=sample_file,
            processing_type="analyze",
        )

        print(f"Created analysis task: {analysis_task.task_id}")

        # Create compression task
        compression_task = ProcessingTask(
            task_id=str(uuid.uuid4()),
            content_file=sample_file,
            processing_type="compress",
            parameters={"compression_level": "medium", "target_format": "mp4"},
        )

        print(f"Created compression task: {compression_task.task_id}")

        # Get processing statistics
        stats = await processor.get_processing_statistics()
        print(f"Processing statistics: {stats}")

        # In a real scenario, would start the processing service
        # await processor.start_processing_service()

    # Run test
    asyncio.run(test_processing_system())
